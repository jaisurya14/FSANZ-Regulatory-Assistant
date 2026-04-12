"""
Generate FSANZ_Project_Proposal.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup: A4, 2.5 cm margins ──────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Cm(21.0)
section.page_height   = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

NAVY  = RGBColor(0x1F, 0x39, 0x64)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'BFBFBF')
        borders.append(el)
    tcPr.append(borders)

def body_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)
    pf.line_spacing = Pt(17)
    return p

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.name      = 'Calibri'
    run.font.size      = Pt(12)
    run.font.bold      = True
    run.font.color.rgb = WHITE
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '1F3964')
    pPr.append(shd)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '120')
    pPr.append(ind)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name      = 'Calibri'
    run.font.size      = Pt(11)
    run.font.bold      = True
    run.font.color.rgb = NAVY
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after  = Pt(2)
    return p

# ── Footer with page number ───────────────────────────────────────────────────
def add_page_number_footer(sec):
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size      = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    fld_start = OxmlElement('w:fldChar')
    fld_start.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.text = ' PAGE '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_start)
    run._r.append(instr)
    run._r.append(fld_end)
    sep = p.add_run('  |  FSANZ Regulatory Affairs Assistant — Project Proposal')
    sep.font.size      = Pt(9)
    sep.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

add_page_number_footer(section)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

t1 = doc.add_paragraph()
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = t1.add_run('FSANZ Regulatory Affairs Assistant')
r1.font.name      = 'Calibri'
r1.font.size      = Pt(22)
r1.font.bold      = True
r1.font.color.rgb = NAVY

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('An AI-Powered Chatbot for Food Standards Compliance')
r2.font.name      = 'Calibri'
r2.font.size      = Pt(15)
r2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
t2.paragraph_format.space_after = Pt(4)

# divider line
div = doc.add_paragraph()
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = div._p.get_or_add_pPr()
pb  = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'),   'single')
bot.set(qn('w:sz'),    '12')
bot.set(qn('w:space'), '1')
bot.set(qn('w:color'), '1F3964')
pb.append(bot)
pPr.append(pb)
div.paragraph_format.space_after = Pt(24)

for _ in range(2):
    doc.add_paragraph()

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('Master of Data Science — Project Team (5 Students)')
r3.font.name      = 'Calibri'
r3.font.size      = Pt(12)
r3.font.bold      = True
r3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run('April 2026  |  12-Week Project  |  480 Person-Hours')
r4.font.name      = 'Calibri'
r4.font.size      = Pt(10)
r4.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PROJECT AIMS AND CONTEXT
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('1. Project Aims and Context')

body_para(
    'Anyone who has ever tried to bring a new food product to market in Australia or New Zealand quickly '
    'discovers that regulatory compliance is anything but straightforward. The FSANZ Food Standards Code '
    'is the definitive legal reference covering food additives, allergen declarations, labelling '
    'requirements, and nutritional standards. It runs to 851 pages of dense regulatory language, '
    'and keeping up with it requires genuine specialist knowledge. For a large manufacturer with an '
    'in-house regulatory affairs team, this is manageable. For a small food business, a startup, or '
    'a product developer working independently, navigating this document is time-consuming, error-prone, '
    'and often requires expensive external advice that many cannot afford.'
)
body_para(
    'The core aim of this project is to remove that barrier. We are building an AI-powered chat '
    'assistant that allows any user, regardless of their regulatory background, to ask plain English '
    'questions about the FSANZ Food Standards Code and receive accurate, clearly cited answers drawn '
    'directly from the regulatory text. Instead of manually searching through hundreds of pages to '
    'determine whether a particular preservative is permitted in a specific food category, a user can '
    'simply type their question and receive a structured answer that references the exact standard '
    'number and page.'
)
body_para(
    'Beyond the conversational interface, we are also developing an AI Compliance Checker module. This '
    'feature allows a user to paste a product ingredient list directly into the system. The assistant '
    'then automatically extracts each ingredient and its declared amount, retrieves the relevant FSANZ '
    'regulation for each one, and delivers a compliance verdict of PASS, WARNING, or FAIL, along with '
    'a practical recommendation wherever an issue is identified. This transforms what would normally '
    'be a multi-hour manual review into an instant, structured report.'
)
body_para(
    'The tool is designed primarily for food product developers, regulatory affairs professionals, '
    'and small to medium food businesses operating under Australian and New Zealand food law. These '
    'groups stand to benefit most from fast, reliable access to regulatory guidance without needing '
    'to commission specialist legal advice for every routine compliance question.'
)
body_para(
    'This project is being delivered by a team of five Master of Data Science students over a '
    'twelve-week period, with approximately 480 total person-hours available across the team. The '
    'scope has been designed to be technically ambitious but achievable within that constraint, '
    'combining real-world cloud infrastructure, state-of-the-art language model technology, and a '
    'practical end-user application that addresses a genuine gap in the market.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('2. Methodology')

body_para(
    'The system is built around a Retrieval Augmented Generation (RAG) architecture, which combines '
    'the precision of information retrieval with the language fluency of a large language model. '
    'The pipeline is divided into four stages, each building on the previous one.'
)

add_h2('Stage 1 — Data Collection and Parsing')
body_para(
    'The primary data source is the FSANZ Food Standards Code PDF, downloaded directly from the '
    'Federal Register of Legislation at legislation.gov.au. This is the authoritative March 2025 '
    'compilation comprising 851 pages. In addition, supplementary guidance pages from the FSANZ '
    'website — covering labelling guidance, allergen declarations, and business FAQs — are collected '
    'using Python\'s Requests library and BeautifulSoup for HTML parsing.'
)
body_para(
    'Text is extracted from the PDF using PyMuPDF (fitz), which reliably handles the multi-column '
    'and table-heavy formatting found throughout the Code. The extracted text is then divided into '
    '800-character chunks with a 100-character overlap using LangChain\'s RecursiveCharacterTextSplitter. '
    'Chunking is essential because large language models cannot process 851 pages in a single prompt. '
    'By breaking the document into small, searchable pieces, the system retrieves only the handful of '
    'sections genuinely relevant to each query, rather than overwhelming the model with irrelevant '
    'content. All raw data, extracted text, and processed chunks are stored in AWS S3.'
)

add_h2('Stage 2 — Knowledge Base and Vector Search')
body_para(
    'Each text chunk is converted into a 384-dimensional numerical vector using SentenceTransformer '
    '(all-MiniLM-L6-v2), a compact and well-performing sentence embedding model. These vectors '
    'capture the semantic meaning of each chunk rather than just its keywords. The vectors are then '
    'uploaded to Pinecone, a managed cloud vector database that supports extremely fast nearest-'
    'neighbour search across tens of thousands of vectors.'
)
body_para(
    'We chose vector search over keyword-based search for a clear practical reason: regulatory '
    'questions are rarely phrased the same way as the legal text they relate to. A user asking '
    '"Can I use potassium sorbate in a fruit drink?" will not match an exact phrase in the FSANZ '
    'Code, but semantic search will correctly identify the relevant standard on permitted '
    'preservatives in beverages. Retrieval quality is evaluated using NDCG@3 against a manually '
    'constructed gold standard set of 20 questions with known correct clause references.'
)

add_h2('Stage 3 — Retrieval Augmented Generation (RAG) Pipeline')
body_para(
    'When a user submits a question, it is first converted to a vector using the same '
    'SentenceTransformer model. Pinecone returns the five most semantically similar chunks from '
    'the FSANZ Code. These chunks, together with the original question, are assembled into a '
    'structured prompt and sent to Anthropic\'s Claude (claude-sonnet-4-6), which generates a '
    'clear, structured answer that includes the relevant standard reference and the page number '
    'of the source text.'
)
body_para(
    'RAG is strongly preferred over using a language model in isolation because it grounds every '
    'answer in current, specific regulatory text. A standalone LLM would rely on knowledge '
    'acquired during training, which may be outdated, incomplete, or insufficiently detailed for '
    'compliance purposes. By retrieving directly from the March 2025 FSANZ Code, the system '
    'ensures answers reflect the actual current rules. Response quality is measured using '
    'BERTScore and ROUGE-L against a gold standard of ten expert-written ideal answers.'
)

add_h2('Stage 4 — AI Compliance Checker')
body_para(
    'The compliance checker extends the RAG pipeline into a structured, multi-step workflow. The '
    'user pastes a product ingredient list into the interface. Claude first parses the text to '
    'extract all individual ingredients and their declared amounts in a structured format. For '
    'each ingredient, the RAG pipeline retrieves the most relevant FSANZ clause, and Claude then '
    'compares the declared amount against the regulatory limit and returns a verdict of PASS, '
    'WARNING, or FAIL. Where an issue is identified, a plain English recommendation is provided. '
    'Every compliance report is automatically logged to AWS S3 with a timestamp, creating an '
    'auditable record of all checks performed.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — HARDWARE, SOFTWARE AND RESOURCES
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('3. Hardware, Software and Resources')

body_para(
    'The project is designed to be low-cost and accessible. Development takes place on standard '
    'student laptops running Python 3.10 or above. Deployment is handled via an AWS EC2 t2.micro '
    'instance, which falls within the AWS Free Tier and incurs no infrastructure cost. The only '
    'meaningful financial outlay is for API usage, estimated at under $20 in total across the '
    'twelve-week project, covering calls to the Anthropic Claude API. All software dependencies '
    'are either open-source or available on free tiers, as shown below. No specialised hardware, '
    'licensed software, or industry-provided infrastructure is required.'
)

# Software table
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
col_w = [Cm(4.8), Cm(7.2), Cm(4.0)]
for i, cell in enumerate(tbl.rows[0].cells):
    cell.width = col_w[i]

hdr_cells = tbl.rows[0].cells
for cell in hdr_cells:
    set_cell_bg(cell, '1F3964')
    set_cell_borders(cell)
for txt, cell in zip(['Tool / Library', 'Purpose', 'Cost'], hdr_cells):
    p   = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name      = 'Calibri'
    run.font.size      = Pt(10)
    run.font.bold      = True
    run.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

sw_rows = [
    ('Python 3.10',          'Core programming language',                      'Free'),
    ('PyMuPDF',              'PDF text extraction',                             'Free'),
    ('LangChain',            'Text chunking pipeline',                          'Free'),
    ('SentenceTransformer',  'Sentence embedding model (all-MiniLM-L6-v2)',    'Free'),
    ('Pinecone',             'Cloud vector database',                           'Free tier'),
    ('Anthropic Claude API', 'LLM answer generation (claude-sonnet-4-6)',      '~$15 est.'),
    ('AWS S3',               'Cloud storage for data and query logs',           'Free tier'),
    ('FastAPI',              'REST API backend server',                         'Free'),
    ('Streamlit',            'Web-based chat user interface',                   'Free'),
    ('boto3',                'Python SDK for AWS services',                     'Free'),
]
for i, (tool, purpose, cost) in enumerate(sw_rows):
    row  = tbl.add_row()
    fill = 'EBF0FA' if i % 2 == 0 else 'FFFFFF'
    for j, (cell, text) in enumerate(zip(row.cells, [tool, purpose, cost])):
        set_cell_bg(cell, fill)
        set_cell_borders(cell)
        p   = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — DATASETS
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('4. Datasets')

body_para(
    'Three datasets are used in this project. All are either publicly available or constructed '
    'internally by the team from public sources. None contain personal data or sensitive '
    'commercial information.'
)

add_h2('Dataset 1 — FSANZ Food Standards Code PDF')
body_para(
    'The primary dataset is the official FSANZ Food Standards Code, downloaded from the Federal '
    'Register of Legislation at legislation.gov.au. The March 2025 compilation is used, which is '
    '851 pages in PDF format. This document is publicly available under Commonwealth legislation '
    'and may be freely used for research and educational purposes without copyright restriction. '
    'It contains no personal data — it is a legal regulatory document.'
)

add_h2('Dataset 2 — FSANZ Website Regulatory Guidance')
body_para(
    'Supplementary guidance material is collected by scraping publicly accessible pages from '
    'foodstandards.gov.au, including labelling guidance notes, allergen declaration guidance, '
    'and business-oriented regulatory FAQs. This material is in HTML format and is scraped using '
    'Python\'s Requests library and BeautifulSoup. All content is publicly available, requires '
    'no authentication, and contains no personal or commercially sensitive data.'
)

add_h2('Dataset 3 — Gold Standard Evaluation Sets (Team-Constructed)')
body_para(
    'To evaluate system performance, the team manually constructs two evaluation sets from public '
    'regulatory text. The retrieval gold standard consists of 20 questions, each paired with the '
    'specific FSANZ clause and page number that constitutes the correct answer, used to calculate '
    'NDCG@3 scores. The response gold standard consists of 10 questions paired with carefully '
    'written ideal answers, used to calculate BERTScore F1 and ROUGE-L scores. These datasets are '
    'not provided by any industry partner; they are constructed entirely internally and are based '
    'exclusively on publicly available regulatory information.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — CHALLENGES AND RISKS
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('5. Challenges and Risks')

body_para(
    'Four risks have been identified and mitigated in the project design.'
)
body_para(
    'The first concerns retrieval quality. If the document is chunked poorly — for example, with '
    'chunks that split a regulation mid-sentence — the system may fail to retrieve the most '
    'relevant clause for a given query. This is mitigated by evaluating retrieval with NDCG@3 '
    'early in the project and adjusting chunk size and overlap parameters if scores fall below '
    'the 0.70 target threshold.'
)
body_para(
    'The second risk is answer hallucination. Large language models can generate confident-sounding '
    'but factually incorrect responses, which would be particularly harmful in a regulatory context. '
    'This is mitigated by strictly constraining the model prompt to answer only from retrieved text, '
    'and by instructing Claude to explicitly state when relevant information was not found rather '
    'than speculating.'
)
body_para(
    'Third, the compliance checker involves multiple sequential AI calls per ingredient, creating '
    'several points of potential failure. This is mitigated by implementing error handling and '
    'fallback responses at each step, and by clearly communicating to users that the tool provides '
    'regulatory guidance rather than legal advice.'
)
body_para(
    'Finally, there is a data coverage risk. Some FSANZ regulations appear inside tables or are '
    'embedded in cross-referenced schedules, which may not be well-represented in plain-text chunks. '
    'This is partially addressed by including supplementary FSANZ website content and by targeting '
    'evaluation questions specifically toward these edge-case areas during testing.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — EXPECTED DELIVERABLES
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('6. Expected Deliverables')

body_para(
    'Six measurable deliverables are defined for this project, each with a clear and testable '
    'success criterion.'
)

# Deliverables table
del_tbl = doc.add_table(rows=1, cols=3)
del_tbl.style = 'Table Grid'
del_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
del_col_w = [Cm(4.0), Cm(7.0), Cm(5.0)]
for i, cell in enumerate(del_tbl.rows[0].cells):
    cell.width = del_col_w[i]

d_hdr = del_tbl.rows[0].cells
for cell in d_hdr:
    set_cell_bg(cell, '1F3964')
    set_cell_borders(cell)
for txt, cell in zip(['Deliverable', 'Description', 'Success Criterion'], d_hdr):
    p   = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name      = 'Calibri'
    run.font.size      = Pt(10)
    run.font.bold      = True
    run.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

del_rows_data = [
    ('Processed Knowledge Base',
     'Pinecone index containing all FSANZ text chunks from the PDF and supplementary website guidance.',
     'NDCG@3 >= 0.70 on the retrieval gold standard.'),
    ('RAG Chat Interface',
     'Streamlit chat application answering regulatory questions with cited standard references and page numbers.',
     'BERTScore F1 >= 0.75; ROUGE-L >= 0.60; 85% accuracy on 10 gold standard questions.'),
    ('AI Compliance Checker',
     'Automated ingredient compliance report with PASS / WARNING / FAIL verdicts and recommendations for each ingredient.',
     '80% classification accuracy across 20 known test cases.'),
    ('AWS Cloud Integration',
     'S3 storage for PDF, chunked data, query logs, and compliance reports with timestamps.',
     'All files stored correctly; every query logged with timestamp and query text.'),
    ('System Documentation',
     'Setup guide, system architecture overview, and API endpoint documentation.',
     'A new team member can run the full system from the documentation alone.'),
    ('Evaluation Report',
     'Written report presenting NDCG@3, BERTScore, ROUGE-L, and compliance accuracy with analysis against targets.',
     'All four metrics reported and discussed with reference to success thresholds.'),
]
for i, (name, desc, criterion) in enumerate(del_rows_data):
    row  = del_tbl.add_row()
    fill = 'EBF0FA' if i % 2 == 0 else 'FFFFFF'
    for j, (cell, text) in enumerate(zip(row.cells, [name, desc, criterion])):
        set_cell_bg(cell, fill)
        set_cell_borders(cell)
        p   = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

doc.add_paragraph()

body_para(
    'The Processed Knowledge Base forms the foundation of the entire system — without a high-quality '
    'retrieval index, every downstream component will underperform, making it the highest-priority '
    'deliverable. The RAG Chat Interface is the primary user-facing feature and is directly assessed '
    'against the response quality gold standard. The AI Compliance Checker extends the system into a '
    'practical workflow tool for product developers. AWS Cloud Integration ensures all data and '
    'activity is persisted reliably and the system can be reproduced by any team member. The '
    'Documentation and Evaluation Report together demonstrate rigour and enable reproducibility '
    'for academic submission.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — PROJECT PLAN
# ═══════════════════════════════════════════════════════════════════════════════
add_h1('7. Project Plan')

body_para(
    'The twelve-week project plan below assigns specific phases and tasks to each two-week period, '
    'with all five team members collaborating during initial setup, final testing, and the '
    'submission week.'
)

plan_tbl = doc.add_table(rows=1, cols=4)
plan_tbl.style = 'Table Grid'
plan_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
plan_col_w = [Cm(1.8), Cm(3.2), Cm(9.2), Cm(1.8)]
for i, cell in enumerate(plan_tbl.rows[0].cells):
    cell.width = plan_col_w[i]

p_hdr = plan_tbl.rows[0].cells
for cell in p_hdr:
    set_cell_bg(cell, '1F3964')
    set_cell_borders(cell)
for txt, cell in zip(['Weeks', 'Phase', 'Tasks', 'Team'], p_hdr):
    p   = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name      = 'Calibri'
    run.font.size      = Pt(10)
    run.font.bold      = True
    run.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

plan_data = [
    ('1 – 2', 'Data Collection & Parsing',
     'Download FSANZ PDF from legislation.gov.au; set up AWS S3 bucket; run web scraper on FSANZ website pages; extract text with PyMuPDF; chunk using LangChain; upload chunks.json to S3.',
     'All'),
    ('3 – 4', 'Knowledge Base',
     'Set up Pinecone index (dimension 384); generate embeddings using SentenceTransformer; upload all chunks to Pinecone; construct retrieval gold standard (20 Q&A pairs); run and review NDCG@3 evaluation.',
     'S1, S2'),
    ('5 – 6', 'RAG Pipeline',
     'Build rag_pipeline.py with vector search and Claude API integration; construct response gold standard (10 Q&A pairs with ideal answers); run BERTScore and ROUGE-L evaluation; refine prompt if needed.',
     'S2, S3'),
    ('7 – 8', 'Backend & API',
     'Build FastAPI server with /health, /ask, and /check-compliance endpoints; develop compliance_checker.py module; connect all endpoints to the RAG pipeline; unit test all API routes.',
     'S3, S4'),
    ('9 – 10', 'Frontend & UI',
     'Build Streamlit interface with Chat and Compliance Checker tabs; implement sidebar with example questions; connect frontend to FastAPI backend; conduct end-to-end integration testing.',
     'S4, S5'),
    ('11', 'Testing',
     'Run full evaluation suite (NDCG@3, BERTScore, ROUGE-L, compliance accuracy); fix any failing tests; validate compliance checker against 20 known test cases; document results.',
     'All'),
    ('12', 'Documentation & Submission',
     'Write setup guide and architecture documentation; prepare final evaluation report; deploy system to AWS EC2 t2.micro; submit all deliverables.',
     'All'),
]
for i, (weeks, phase, tasks, resp) in enumerate(plan_data):
    row  = plan_tbl.add_row()
    fill = 'EBF0FA' if i % 2 == 0 else 'FFFFFF'
    for j, (cell, text) in enumerate(zip(row.cells, [weeks, phase, tasks, resp])):
        set_cell_bg(cell, fill)
        set_cell_borders(cell)
        p   = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

doc.add_paragraph()

# ── Save ─────────────────────────────────────────────────────────────────────
out = r'C:\Users\ganes\downloads\FSANZ-Assistant\FSANZ_Project_Proposal.docx'
doc.save(out)
print(f'Saved: {out}')
